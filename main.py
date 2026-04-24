#!/usr/bin/env python3
"""
Google Maps Scraper - Main Entry Point
A robust scraper for extracting business information from Google Maps using Playwright
"""

import argparse
import asyncio
import logging
import os
import re
from datetime import datetime
from dotenv import load_dotenv
from scraper import GoogleMapsScraper


def setup_output_directory(output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def generate_filename(base_name, extension):
    """Generate filename with timestamp"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base_name}_{timestamp}.{extension}"

def _env_bool(value: str | None, default: bool = False) -> bool:
    if value is None:
        return default
    v = str(value).strip().lower()
    if v in ["1", "true", "yes", "y", "on"]:
        return True
    if v in ["0", "false", "no", "n", "off"]:
        return False
    return default


def build_summary(query: str, results: list[dict]) -> str:
    lines: list[str] = []
    lines.append("SCRAPING SUMMARY")
    lines.append(f"QUERY: {query}")
    lines.append(f"TIMESTAMP: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"TOTAL_ROWS: {len(results)}")

    def non_empty(key: str) -> int:
        return sum(1 for r in results if str(r.get(key, "") or "").strip())

    lines.append(f"WITH_PHONE: {non_empty('phone')}")
    lines.append(f"WITH_EMAIL: {non_empty('email')}")
    lines.append(f"WITH_WEBSITE: {non_empty('website')}")

    def freq(key: str) -> list[tuple[str, int]]:
        counts: dict[str, int] = {}
        for r in results:
            v = str(r.get(key, "") or "").strip()
            if not v:
                continue
            counts[v] = counts.get(v, 0) + 1
        return sorted(counts.items(), key=lambda x: (-x[1], x[0]))

    facility = freq("facility_category")
    if facility:
        lines.append("")
        lines.append("FACILITY_CATEGORY_COUNTS:")
        for k, c in facility:
            lines.append(f"- {k}: {c}")

    sports = freq("location_category")
    if sports:
        lines.append("")
        lines.append("LOCATION_CATEGORY_COUNTS:")
        for k, c in sports[:20]:
            lines.append(f"- {k}: {c}")

    def to_float(x):
        try:
            return float(str(x).strip())
        except Exception:
            return None

    scored: list[tuple[float, int, str]] = []
    for r in results:
        rating = to_float(r.get("rating", ""))
        if rating is None:
            continue
        rc = r.get("reviews_count", "")
        try:
            reviews = int(str(rc).replace(".", "").replace(",", "").strip() or "0")
        except Exception:
            reviews = 0
        scored.append((rating, reviews, str(r.get("name", "") or "").strip()))

    if scored:
        scored.sort(key=lambda x: (-x[0], -x[1], x[2]))
        lines.append("")
        lines.append("TOP_BY_RATING:")
        for rating, reviews, name in scored[:10]:
            lines.append(f"- {name} | rating={rating} | reviews={reviews}")

    return "\n".join(lines) + "\n"


def save_summary_docx(*, summary_text: str, path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in (summary_text or "").splitlines():
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.strip().upper() == "SCRAPING SUMMARY":
            doc.add_heading("Scrapper BOT AI Professional — Summary", level=1)
            continue
        doc.add_paragraph(line)

    doc.save(path)


async def run_scrape_and_exports(
    *,
    query: str,
    headless: bool,
    max_results: int,
    delay_range: tuple[float, float],
    output_dir: str,
    output_format: str,
    ai_analyze: bool,
    ai_categorize: bool,
) -> dict:
    scraper = GoogleMapsScraper(
        headless=headless,
        delay_range=delay_range,
        max_results=max_results
    )

    try:
        results = await scraper.scrape(query)
        if not results:
            return {"results": [], "paths": {}}

        if ai_categorize:
            results = scraper.categorize_location_and_facility_with_ai(results)

        paths: dict[str, str] = {}

        if output_format in ["csv", "both"]:
            csv_filename = generate_filename("google_maps_results", "csv")
            csv_path = os.path.join(output_dir, csv_filename)
            scraper.save_to_csv(results, csv_path)
            paths["csv"] = csv_path

        if output_format in ["excel", "both"]:
            excel_filename = generate_filename("google_maps_results", "xlsx")
            excel_path = os.path.join(output_dir, excel_filename)
            scraper.save_to_excel(results, excel_path)
            paths["xlsx"] = excel_path

        if ai_analyze:
            analysis_result = scraper.analyze_with_ai(results)
            analysis_filename = generate_filename("ai_analysis", "txt")
            analysis_path = os.path.join(output_dir, analysis_filename)
            with open(analysis_path, "w", encoding="utf-8") as f:
                f.write(analysis_result)
            paths["ai_analysis"] = analysis_path

        summary_filename = generate_filename("summary", "docx")
        summary_path = os.path.join(output_dir, summary_filename)
        summary_content = build_summary(query, results)
        save_summary_docx(summary_text=summary_content, path=summary_path)
        paths["summary_docx"] = summary_path

        return {"results": results, "paths": paths}
    finally:
        await scraper.close()


def _parse_kv_options(text: str) -> tuple[str, dict]:
    raw = (text or "").strip()
    if not raw:
        return "", {}

    parts = [p.strip() for p in raw.split("|") if p.strip()]
    query = parts[0] if parts else ""
    opts: dict[str, str] = {}
    for p in parts[1:]:
        if "=" not in p:
            continue
        k, v = p.split("=", 1)
        k = k.strip().lower()
        v = v.strip()
        if k and v:
            opts[k] = v
    return query, opts


async def run_telegram_bot() -> None:
    try:
        from telegram import Update
        from telegram.constants import ChatAction
        from telegram.ext import ApplicationBuilder, CommandHandler, ContextTypes, MessageHandler, filters
        from telegram.error import Conflict
    except Exception as e:
        raise RuntimeError("python-telegram-bot belum terpasang. Jalankan: pip install -r requirements.txt") from e

    token = os.getenv("TELEGRAM_BOT_TOKEN", "").strip()
    if not token:
        raise EnvironmentError("TELEGRAM_BOT_TOKEN tidak ditemukan di environment variables.")

    logging.basicConfig(
        level=logging.WARNING,
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    )
    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("telegram").setLevel(logging.WARNING)

    allowed = os.getenv("TELEGRAM_ALLOWED_CHAT_IDS", "").strip()
    allowed_ids: set[int] = set()
    if allowed:
        for x in re.split(r"[,\s]+", allowed):
            x = x.strip()
            if not x:
                continue
            try:
                allowed_ids.add(int(x))
            except Exception:
                continue

    output_dir = setup_output_directory(os.getenv("OUTPUT_DIRECTORY", "output"))
    delay_min = float(os.getenv("DEFAULT_DELAY_MIN", "1"))
    delay_max = float(os.getenv("DEFAULT_DELAY_MAX", "3"))
    default_delay_range = (delay_min, delay_max)
    default_headless = _env_bool(os.getenv("HEADLESS"), True)
    default_max_results = int(os.getenv("DEFAULT_MAX_RESULTS", "50"))
    default_ai_categorize = _env_bool(os.getenv("AI_CATEGORIZE"), False)

    usage_text = (
        "Scrapper BOT AI Professional\n\n"
        "Cara pakai:\n"
        "1) Ketik /scrape lalu diikuti query.\n"
        "2) Opsi tambahan pakai pemisah | (pipe).\n\n"
        "Format:\n"
        "/scrape <query> | max=50 | headless=1 | ai=1\n\n"
        "Contoh:\n"
        "/scrape Universitas di Jakarta Selatan | max=10 | ai=1\n\n"
        "Catatan:\n"
        "- Output yang dikirim: XLSX + Summary (DOCX)\n"
        "- CSV tidak disediakan di mode Telegram.\n"
    )

    async def _authorized(update: Update) -> bool:
        if not allowed_ids:
            return True
        chat_id = update.effective_chat.id if update.effective_chat else 0
        return chat_id in allowed_ids

    async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not await _authorized(update):
            await update.message.reply_text("Unauthorized.")
            return
        await update.message.reply_text(usage_text)

    async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not await _authorized(update):
            await update.message.reply_text("Unauthorized.")
            return
        await update.message.reply_text(usage_text)

    async def non_command_text(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not await _authorized(update):
            await update.message.reply_text("Unauthorized.")
            return
        await update.message.reply_text("Gunakan perintah /scrape.\n\n" + usage_text)

    async def scrape_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
        if not await _authorized(update):
            await update.message.reply_text("Unauthorized.")
            return

        chat_id = update.effective_chat.id if update.effective_chat else 0

        raw = update.message.text or ""
        rest = raw.split(" ", 1)[1] if " " in raw else ""
        query, opts = _parse_kv_options(rest)
        if not query:
            await update.message.reply_text("Format salah.\n\n" + usage_text)
            return

        allowed_opts = {"max", "headless", "ai"}
        unknown = [k for k in opts.keys() if k not in allowed_opts]
        if unknown:
            await update.message.reply_text(
                "Opsi tidak dikenali: " + ", ".join(sorted(unknown)) + "\n\n" + usage_text
            )
            return

        if len(query.strip()) < 3:
            await update.message.reply_text("Query terlalu pendek.\n\n" + usage_text)
            return

        max_results = default_max_results
        if "max" in opts:
            try:
                max_results = int(opts["max"])
            except Exception:
                pass

        headless = default_headless
        if "headless" in opts:
            headless = _env_bool(opts["headless"], default_headless)

        ai_categorize = default_ai_categorize
        if "ai" in opts:
            ai_categorize = _env_bool(opts["ai"], default_ai_categorize)

        await update.message.reply_text(
            "Scrapper BOT AI Professional — Proses dimulai\n"
            f"- QUERY: {query}\n"
            f"- MAX_RESULTS: {max_results}\n"
            f"- HEADLESS: {'YES' if headless else 'NO'}\n"
            f"- AI_CATEGORIZE: {'YES' if ai_categorize else 'NO'}\n"
            "- OUTPUT: XLSX + SUMMARY(DOCX)"
        )

        typing_stop = asyncio.Event()

        async def _typing_loop() -> None:
            while not typing_stop.is_set():
                try:
                    await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
                except Exception:
                    pass
                await asyncio.sleep(4)

        typing_task = asyncio.create_task(_typing_loop())

        out: dict | None = None
        try:
            try:
                out = await run_scrape_and_exports(
                    query=query,
                    headless=headless,
                    max_results=max_results,
                    delay_range=default_delay_range,
                    output_dir=output_dir,
                    output_format="excel",
                    ai_analyze=False,
                    ai_categorize=ai_categorize,
                )
            except OSError as e:
                if getattr(e, "errno", None) == 28:
                    await update.message.reply_text("Gagal menyimpan file: disk penuh (No space left on device).")
                    return
                raise
            except Exception as e:
                await update.message.reply_text(f"Gagal scraping: {e}")
                return
        finally:
            typing_stop.set()
            try:
                typing_task.cancel()
            except Exception:
                pass

        if out is None:
            await update.message.reply_text("Gagal scraping: output kosong.")
            return

        results = out.get("results") or []
        paths = out.get("paths") or {}
        if not results:
            await update.message.reply_text("Tidak ada hasil.")
            return

        await update.message.reply_text(f"Selesai. Total data: {len(results)}\nFile sedang dikirim...")

        try:
            await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.UPLOAD_DOCUMENT)
        except Exception:
            pass

        for key in ["xlsx", "summary_docx"]:
            p = paths.get(key)
            if not p:
                continue
            try:
                with open(p, "rb") as f:
                    await context.bot.send_document(chat_id=chat_id, document=f, filename=os.path.basename(p))
            except Exception:
                continue

    app = ApplicationBuilder().token(token).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_cmd))
    app.add_handler(CommandHandler("scrape", scrape_cmd))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, non_command_text))
    await app.initialize()
    await app.start()

    if app.updater is None:
        raise RuntimeError("Telegram Updater tidak tersedia untuk mode polling.")

    attempts = 0
    while True:
        try:
            await app.bot.delete_webhook(drop_pending_updates=True)
        except Exception:
            pass
        try:
            await app.updater.start_polling(drop_pending_updates=True)
            break
        except Conflict:
            attempts += 1
            if attempts >= 5:
                raise
            await asyncio.sleep(3)

    stop_event = asyncio.Event()
    polling_restart_attempts = 0
    polling_error: dict[str, str] = {}

    async def _ensure_polling_running() -> None:
        nonlocal polling_restart_attempts
        while not stop_event.is_set():
            await asyncio.sleep(2)
            if app.updater is None:
                polling_error["reason"] = "Updater tidak tersedia."
                stop_event.set()
                return
            if app.updater.running:
                continue
            polling_restart_attempts += 1
            try:
                await app.bot.delete_webhook(drop_pending_updates=True)
            except Exception:
                pass
            try:
                await app.updater.start_polling(drop_pending_updates=True)
            except Conflict:
                if polling_restart_attempts >= 5:
                    polling_error["reason"] = (
                        "Conflict: ada instance bot lain yang sedang polling atau webhook masih aktif. "
                        "Pastikan hanya satu proses bot berjalan dan webhook sudah dihapus."
                    )
                    stop_event.set()
                    return
                await asyncio.sleep(3)
            except Exception as e:
                if polling_restart_attempts >= 5:
                    polling_error["reason"] = f"Polling gagal berulang: {e}"
                    stop_event.set()
                    return
                await asyncio.sleep(3)
    monitor_task: asyncio.Task | None = None
    try:
        import signal

        loop = asyncio.get_running_loop()
        for sig in (signal.SIGINT, signal.SIGTERM):
            try:
                loop.add_signal_handler(sig, stop_event.set)
            except NotImplementedError:
                pass

        print("🤖 Telegram bot berjalan. Tekan Ctrl+C untuk berhenti.")
        monitor_task = asyncio.create_task(_ensure_polling_running())
        await stop_event.wait()
    finally:
        if monitor_task is not None:
            try:
                monitor_task.cancel()
            except Exception:
                pass
        try:
            await app.updater.stop()
        except Exception:
            pass
        try:
            await app.stop()
        except Exception:
            pass
        try:
            await app.shutdown()
        except Exception:
            pass
        if polling_error.get("reason"):
            raise RuntimeError(polling_error["reason"])


async def main():
    parser = argparse.ArgumentParser(
        description="Google Maps Scraper - Extract business information from Google Maps"
    )
    parser.add_argument(
        "query",
        type=str,
        nargs="?",
        default=None,
        help="Search query (e.g., 'restaurants in New York'). Tidak diperlukan jika menjalankan --telegram-bot"
    )
    parser.add_argument(
        "--headless",
        action="store_true",
        default=False,
        help="Run browser in headless mode"
    )
    parser.add_argument(
        "--max-results",
        type=int,
        default=50,
        help="Maximum number of results to scrape (default: 50)"
    )
    parser.add_argument(
        "--delay",
        type=float,
        default=None,
        help="Delay antar interaksi (detik). Jika diisi, akan menjadi delay tetap (tanpa random)."
    )
    parser.add_argument(
        "--output-format",
        type=str,
        choices=["csv", "excel", "both"],
        default="both",
        help="Output format: csv, excel, or both (default: both)"
    )
    parser.add_argument(
        "--ai-analyze",
        action="store_true",
        help="Analyze the scraped results using Groq AI"
    )
    parser.add_argument(
        "--ai-categorize",
        action="store_true",
        help="Gunakan AI untuk mengisi LOCATION_CATEGORY dan FACILITY_CATEGORY"
    )
    parser.add_argument(
        "--telegram-bot",
        action="store_true",
        help="Jalankan bot Telegram (polling). Gunakan TELEGRAM_BOT_TOKEN di env."
    )

    args = parser.parse_args()

    # Load environment variables (e.g. GROQ_API_KEY)
    load_dotenv()

    if args.telegram_bot:
        await run_telegram_bot()
        return
    if not args.query:
        parser.error("the following arguments are required: query (atau gunakan --telegram-bot)")

    # Setup output directory
    output_dir = setup_output_directory(os.getenv("OUTPUT_DIRECTORY", "output"))

    delay_min = float(os.getenv("DEFAULT_DELAY_MIN", "1"))
    delay_max = float(os.getenv("DEFAULT_DELAY_MAX", "3"))
    delay_range = (delay_min, delay_max)
    if args.delay is not None:
        delay_range = (args.delay, args.delay)

    # Initialize scraper
    scraper = GoogleMapsScraper(
        headless=args.headless,
        delay_range=delay_range,
        max_results=args.max_results
    )

    try:
        print(f"🔍 Starting Google Maps scraper for query: '{args.query}'")
        print(f"📊 Maximum results: {args.max_results}")
        print(f"🌐 Headless mode: {'Yes' if args.headless else 'No'}")
        print("=" * 50)

        # Run scraper
        results = await scraper.scrape(args.query)

        if not results:
            print("❌ No results found or scraping failed")
            return

        if args.ai_categorize or os.getenv("AI_CATEGORIZE", "").lower() in ["1", "true", "yes", "y"]:
            print("🤖 Running AI categorization for LOCATION_CATEGORY and FACILITY_CATEGORY...")
            results = scraper.categorize_location_and_facility_with_ai(results)

        print(f"✅ Successfully scraped {len(results)} businesses")

        # Generate output files
        if args.output_format in ["csv", "both"]:
            csv_filename = generate_filename("google_maps_results", "csv")
            csv_path = os.path.join(output_dir, csv_filename)
            try:
                scraper.save_to_csv(results, csv_path)
                print(f"📁 CSV saved: {csv_path}")
            except OSError as e:
                if getattr(e, "errno", None) == 28:
                    print(f"❌ Gagal menyimpan CSV (disk penuh): {csv_path}")
                else:
                    raise

        if args.output_format in ["excel", "both"]:
            excel_filename = generate_filename("google_maps_results", "xlsx")
            excel_path = os.path.join(output_dir, excel_filename)
            try:
                scraper.save_to_excel(results, excel_path)
                print(f"📁 Excel saved: {excel_path}")
            except OSError as e:
                if getattr(e, "errno", None) == 28:
                    print(f"❌ Gagal menyimpan Excel (disk penuh): {excel_path}")
                else:
                    raise

        # Run AI Analysis if requested
        if args.ai_analyze:
            print("\n" + "=" * 50)
            analysis_result = scraper.analyze_with_ai(results)
            print("=" * 50)
            print("🤖 AI Analysis Report:\n")
            print(analysis_result)
            print("\n" + "=" * 50)

            # Save analysis to file
            analysis_filename = generate_filename("ai_analysis", "txt")
            analysis_path = os.path.join(output_dir, analysis_filename)
            try:
                with open(analysis_path, "w", encoding="utf-8") as f:
                    f.write(analysis_result)
                print(f"📁 AI Analysis saved: {analysis_path}")
            except OSError as e:
                if getattr(e, "errno", None) == 28:
                    print(f"❌ Gagal menyimpan AI analysis (disk penuh): {analysis_path}")
                else:
                    raise

        summary_filename = generate_filename("summary", "txt")
        summary_path = os.path.join(output_dir, summary_filename)
        summary_content = build_summary(args.query, results)
        try:
            with open(summary_path, "w", encoding="utf-8") as f:
                f.write(summary_content)
            print(f"📁 Summary saved: {summary_path}")
        except OSError as e:
            if getattr(e, "errno", None) == 28:
                print(f"❌ Gagal menyimpan summary (disk penuh): {summary_path}")
            else:
                raise

        print("🎉 Scraping completed successfully!")

    except KeyboardInterrupt:
        print("\n⚠️  Scraping interrupted by user")
    except Exception as e:
        print(f"❌ Error during scraping: {str(e)}")
    finally:
        await scraper.close()


if __name__ == "__main__":
    asyncio.run(main())
